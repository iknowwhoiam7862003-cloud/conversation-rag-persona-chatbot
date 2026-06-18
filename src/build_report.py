from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .utils import DATA_DIR, REPORT_DIR, read_json


def _set_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_report(output_path: Path = REPORT_DIR / "RAG_Project_Documentation.docx") -> Path:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    parse_stats = read_json(DATA_DIR / "parse_stats.json")
    checkpoint_stats = read_json(DATA_DIR / "checkpoint_stats.json")
    index_stats = read_json(Path(__file__).resolve().parents[1] / "indexes" / "index_stats.json")

    doc = Document()
    _set_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Conversation RAG System with Persona Chatbot")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Project documentation for chronological checkpointing, local retrieval, persona extraction, and chatbot querying.").italic = True

    doc.add_heading("Project Objective", level=1)
    doc.add_paragraph(
        "The project turns a CSV of daily conversations into a local-first intelligent system. "
        "It preserves chronological order, detects topic changes over time, creates independent "
        "100-message summaries, extracts evidence-backed personas for both speakers, and exposes "
        "a chatbot that can answer questions using retrieved summaries and raw message chunks."
    )

    doc.add_heading("Dataset Summary", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    for metric, value in [
        ("CSV rows / conversation days", parse_stats["rows"]),
        ("Parsed messages", parse_stats["messages"]),
        ("User 1 messages", parse_stats["User 1"]),
        ("User 2 messages", parse_stats["User 2"]),
        ("Topic checkpoints", checkpoint_stats["topic_checkpoints"]),
        ("100-message checkpoints", checkpoint_stats["msg100_checkpoints"]),
        ("Message chunks indexed", index_stats["chunk_docs"]),
    ]:
        row = table.add_row().cells
        row[0].text = str(metric)
        row[1].text = str(value)

    doc.add_heading("Architecture", level=1)
    _add_bullets(
        doc,
        [
            "Parser: reads the single-column CSV and flattens rows into globally ordered message records.",
            "Checkpoint engine: creates semantic topic checkpoints and fixed 100-message checkpoints independently.",
            "Local summarizer: selects high-signal chronological messages using TF-IDF-style keyword scoring.",
            "Index store: builds searchable local TF-IDF indexes for topics, fixed checkpoints, and message chunks.",
            "Persona extractor: creates separate JSON personas for User 1 and User 2 with message-level evidence.",
            "Chatbot: routes persona questions to persona JSON and conversation questions to RAG retrieval.",
        ],
    )

    doc.add_heading("How Topic Checkpoints Work", level=1)
    doc.add_paragraph(
        "The system streams messages in chronological order. Every five messages, once a minimum topic "
        "length is reached, it compares the current message against a rolling window of the previous "
        "15 messages. A low cosine similarity score indicates semantic drift, so the previous segment "
        "is closed as a topic checkpoint and summarized independently."
    )

    doc.add_heading("How 100-Message Checkpoints Work", level=1)
    doc.add_paragraph(
        "The 100-message checkpoints are independent of topic boundaries. Every block of 100 chronological "
        "messages receives its own summary and keywords. These checkpoints provide broad coverage even when "
        "topic segmentation is imperfect."
    )

    doc.add_heading("Query Handling", level=1)
    _add_bullets(
        doc,
        [
            "The query is searched against topic summaries, 100-message summaries, and raw message chunks.",
            "The response includes checkpoint IDs, message ranges, scores, summaries, and raw evidence previews.",
            "Persona-style questions are answered from persona JSON and supplemented with RAG evidence.",
        ],
    )

    doc.add_heading("Persona Extraction", level=1)
    doc.add_paragraph(
        "Persona extraction is evidence-first. Habits, facts, life events, and traits are stored only when "
        "the system finds supporting conversation signals. Quantitative communication style metrics are "
        "computed directly from messages, including average words per message, question ratio, exclamation "
        "ratio, emoji usage, common openings, and tone markers."
    )

    doc.add_heading("Tools and Libraries Used", level=1)
    _add_bullets(
        doc,
        [
            "Python standard library for parsing, JSONL persistence, checkpointing, local TF-IDF, and CLI scripts.",
            "python-docx for generating this Word documentation.",
            "Streamlit for the chatbot UI.",
            "Optional: scikit-learn, sentence-transformers, and FAISS can be installed for stronger retrieval later.",
        ],
    )

    doc.add_heading("Test Cases", level=1)
    _add_bullets(
        doc,
        [
            "Parser count check: verify row count, total messages, and per-speaker messages.",
            "Checkpoint coverage check: verify topic ranges are chronological and non-overlapping.",
            "100-message checkpoint check: verify full checkpoints contain exactly 100 messages except the final partial block.",
            "Retrieval smoke tests: ask about habits, communication style, personal facts, and common topics.",
            "Persona evidence check: confirm qualitative claims include message IDs.",
            "Chatbot UI check: run Streamlit and test the required sample questions.",
        ],
    )

    doc.add_heading("Limitations and Future Improvements", level=1)
    _add_bullets(
        doc,
        [
            "The default summaries are extractive, so they are reliable but less fluent than LLM summaries.",
            "The TF-IDF topic detector is fast and local, but threshold tuning affects checkpoint granularity.",
            "The system treats User 1 and User 2 as stable identities across all rows because that was selected for this task.",
            "A future version can add optional local transformer embeddings or an API-backed rewrite step without changing the data pipeline.",
        ],
    )

    doc.add_heading("How to Run", level=1)
    for command in [
        r"python -m src.parse --input C:\Users\syedf\Downloads\conversations.csv",
        "python -m src.checkpoints",
        "python -m src.index_store",
        "python -m src.persona",
        'python -m src.chatbot "What kind of person is User 1?"',
        "streamlit run app.py",
    ]:
        para = doc.add_paragraph()
        para.style = "Intense Quote"
        para.add_run(command)

    doc.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Build DOCX project documentation.")
    parser.add_argument("--output", default=str(REPORT_DIR / "RAG_Project_Documentation.docx"))
    args = parser.parse_args()
    print(build_report(Path(args.output)))


if __name__ == "__main__":
    main()
